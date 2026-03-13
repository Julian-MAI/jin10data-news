#!/usr/bin/env python3
"""
Jin10 news collector.

Features:
- Fetch latest Jin10 flash/news items from public endpoint.
- Optional pagination by max pages and delay.
- Optional keyword filtering.
- Auto-generate two Chinese Word reports only:
    1) yesterday summary
    2) today's latest summary

Usage examples:
  python jin10_news_collector.py
  python jin10_news_collector.py --pages 5 --keyword 美联储 --output-dir output
"""

from __future__ import annotations

import argparse
import html
import json
import os
import re
import time
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

API_URL = "https://flash-api.jin10.com/get_flash_list"
DEFAULT_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    # Jin10 endpoint commonly validates these request headers.
    "x-app-id": "bVBF4FyRTn5NJF5n",
    "x-version": "1.0.0",
    "Referer": "https://www.jin10.com/",
    "Origin": "https://www.jin10.com",
    "Accept": "application/json, text/plain, */*",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
}

CATEGORY_KEYWORDS = {
    "宏观政策": [
        "美联储", "降息", "加息", "通胀", "cpi", "ppi", "非农", "失业率", "gdp", "央行", "国债", "财政", "货币政策",
    ],
    "外汇贵金属": [
        "美元", "欧元", "英镑", "日元", "汇率", "外汇", "黄金", "白银", "现货金", "伦敦金", "美债收益率",
    ],
    "能源与大宗": [
        "原油", "wti", "布伦特", "天然气", "opec", "铜", "铝", "镍", "铁矿石", "大豆", "玉米", "大宗商品",
    ],
    "股市公司": [
        "a股", "港股", "美股", "纳指", "标普", "道指", "财报", "业绩", "回购", "分红", "特斯拉", "英伟达", "苹果",
    ],
    "加密市场": [
        "比特币", "btc", "以太坊", "eth", "加密", "稳定币", "区块链", "交易所", "链上",
    ],
}


def parse_jin10_time(ts: str) -> Optional[datetime]:
    if not ts:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M"):
        try:
            return datetime.strptime(ts[:19], fmt)
        except ValueError:
            continue
    return None


def strip_html(text: str) -> str:
    if not text:
        return ""
    cleaned = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    cleaned = html.unescape(cleaned)
    cleaned = cleaned.replace("\u3000", " ")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    # Remove common trailing noise from flash/news feeds.
    noise_patterns = [
        r"（?金十数据.*?）?$",
        r"点击查看.*$",
        r"详见.*?APP.*$",
        r"来源[:：].*$",
    ]
    for pattern in noise_patterns:
        cleaned = re.sub(pattern, "", cleaned).strip()

    return cleaned


def normalize_item(item: Dict[str, Any]) -> Dict[str, Any]:
    # Different endpoint versions may return fields under either root or nested data object.
    data_obj = item.get("data", {}) if isinstance(item.get("data"), dict) else {}

    news_id = item.get("id") or data_obj.get("id") or ""
    timestamp = item.get("time") or data_obj.get("time") or ""
    text = item.get("content") or data_obj.get("content") or item.get("title") or ""
    text = strip_html(text)

    country = item.get("country") or data_obj.get("country") or ""
    importance = item.get("important")
    if importance is None:
        importance = data_obj.get("important", "")
    channel = item.get("channel") or data_obj.get("channel") or ""

    parsed_dt = parse_jin10_time(timestamp)

    return {
        "id": str(news_id),
        "time": timestamp,
        "datetime": parsed_dt.strftime("%Y-%m-%d %H:%M:%S") if parsed_dt else "",
        "content": text,
        "country": country,
        "importance": importance,
        "channel": channel,
    }


def request_page(max_time: Optional[str], timeout: int) -> Dict[str, Any]:
    params: Dict[str, str] = {}
    if max_time:
        params["max_time"] = max_time

    resp = requests.get(API_URL, headers=DEFAULT_HEADERS, params=params, timeout=timeout)
    resp.raise_for_status()

    payload = resp.json()
    if not isinstance(payload, dict):
        raise ValueError("Unexpected response format (not a JSON object)")

    return payload


def extract_items(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    # Compatible with multiple field names returned by different versions.
    candidates = [
        payload.get("data"),
        payload.get("list"),
        payload.get("news"),
    ]

    for node in candidates:
        if isinstance(node, list):
            return [x for x in node if isinstance(x, dict)]
        if isinstance(node, dict):
            for key in ("data", "list", "news"):
                sub = node.get(key)
                if isinstance(sub, list):
                    return [x for x in sub if isinstance(x, dict)]

    return []


def collect_news(
    pages: int,
    timeout: int,
    delay: float,
    since_hours: Optional[int],
    keyword: Optional[str],
) -> List[Dict[str, Any]]:
    all_items: List[Dict[str, Any]] = []
    seen_ids = set()
    max_time: Optional[str] = None

    cutoff = None
    if since_hours is not None:
        cutoff = datetime.now() - timedelta(hours=since_hours)

    for page in range(1, pages + 1):
        payload = request_page(max_time=max_time, timeout=timeout)
        raw_items = extract_items(payload)
        if not raw_items:
            print(f"[INFO] 第 {page} 页无数据，停止翻页。")
            break

        normalized_page = []
        for raw in raw_items:
            item = normalize_item(raw)
            if not item["content"]:
                continue

            # Deduplicate by id first, then fallback by timestamp+content hash key.
            dedup_key = item["id"] or f"{item['time']}|{item['content'][:60]}"
            if dedup_key in seen_ids:
                continue
            seen_ids.add(dedup_key)

            if cutoff and item["datetime"]:
                dt = datetime.strptime(item["datetime"], "%Y-%m-%d %H:%M:%S")
                if dt < cutoff:
                    continue

            if keyword and keyword not in item["content"]:
                continue

            normalized_page.append(item)

        all_items.extend(normalized_page)

        # Locate earliest time in the current raw page as next pagination anchor.
        times = [normalize_item(x).get("time", "") for x in raw_items]
        times = [t for t in times if t]
        if times:
            max_time = min(times)
        else:
            break

        print(f"[INFO] 已抓取第 {page} 页，新增 {len(normalized_page)} 条，总计 {len(all_items)} 条")
        time.sleep(max(delay, 0.0))

    return all_items


def collect_news_for_date(
    target_date,
    pages: int,
    timeout: int,
    delay: float,
    keyword: Optional[str],
) -> List[Dict[str, Any]]:
    """Collect Jin10 news for a specific date by anchoring max_time to that day 23:59:59."""
    all_items: List[Dict[str, Any]] = []
    seen_ids = set()
    max_time = f"{target_date.strftime('%Y-%m-%d')} 23:59:59"

    for page in range(1, pages + 1):
        payload = request_page(max_time=max_time, timeout=timeout)
        raw_items = extract_items(payload)
        if not raw_items:
            print(f"[INFO] 目标日期抓取第 {page} 页无数据，停止翻页。")
            break

        page_datetimes: List[datetime] = []
        page_added = 0

        for raw in raw_items:
            item = normalize_item(raw)
            if not item["content"]:
                continue

            dt = parse_jin10_time(item.get("time", ""))
            if not dt:
                continue

            page_datetimes.append(dt)

            if dt.date() != target_date:
                continue

            if keyword and keyword not in item["content"]:
                continue

            dedup_key = item["id"] or f"{item['time']}|{item['content'][:60]}"
            if dedup_key in seen_ids:
                continue

            seen_ids.add(dedup_key)
            all_items.append(item)
            page_added += 1

        if not page_datetimes:
            break

        min_dt = min(page_datetimes)
        max_time = min_dt.strftime("%Y-%m-%d %H:%M:%S")

        print(f"[INFO] 目标日期抓取第 {page} 页，新增 {page_added} 条，总计 {len(all_items)} 条")

        if min_dt.date() < target_date:
            # Reached older date; target-date scan is sufficient.
            break

        time.sleep(max(delay, 0.0))

    return all_items


def classify_news(content: str) -> str:
    text = (content or "").lower()
    best_category = "其他"
    best_score = 0

    for category, keywords in CATEGORY_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw.lower() in text)
        if score > best_score:
            best_score = score
            best_category = category

    return best_category


def extract_key_points(items: List[Dict[str, Any]], max_points: int = 5) -> List[str]:
    """从新闻列表中提取关键要点"""
    if not items:
        return []

    # 按重要性排序，importance=1 优先
    sorted_items = sorted(items, key=lambda x: (
        -int(x.get("importance", 0) or 0),
        x.get("datetime", "")
    ), reverse=True)

    key_points = []
    for item in sorted_items[:max_points]:
        content = item.get("content", "").strip()
        if content and len(content) > 10:
            # 清理内容，保留核心信息
            content = " ".join(content.split())
            key_points.append(content)

    return key_points


def generate_category_insight(category: str, items: List[Dict[str, Any]]) -> str:
    """根据分类和新闻内容生成观点分析"""
    if not items:
        return ""

    all_text = " ".join([item.get("content", "") for item in items]).lower()

    insights = {
        "宏观政策": [],
        "外汇贵金属": [],
        "能源与大宗": [],
        "股市公司": [],
        "加密市场": [],
        "其他": [],
    }

    # 宏观政策观点
    if category == "宏观政策":
        if "加息" in all_text or "rate hike" in all_text:
            insights["宏观政策"].append("央行加息预期升温，需关注对风险资产的压制")
        if "降息" in all_text or "rate cut" in all_text:
            insights["宏观政策"].append("降息预期增强，或利好股市和债市")
        if "通胀" in all_text or "cpi" in all_text or "ppi" in all_text:
            insights["宏观政策"].append("通胀数据成为市场焦点，影响货币政策走向")
        if "非农" in all_text or "就业" in all_text or "失业" in all_text:
            insights["宏观政策"].append("就业市场数据影响美联储决策，需持续跟踪")
        if "gdp" in all_text or "经济增长" in all_text:
            insights["宏观政策"].append("经济增长数据反映基本面，影响市场情绪")
        if "关税" in all_text or "tariff" in all_text or "贸易" in all_text:
            insights["宏观政策"].append("贸易政策变化将影响相关行业及汇率走势")
        if not insights["宏观政策"]:
            insights["宏观政策"].append("宏观环境整体稳定，建议关注后续政策动向")

    # 外汇贵金属观点
    elif category == "外汇贵金属":
        if "美元" in all_text and ("走强" in all_text or "上涨" in all_text or "涨" in all_text):
            insights["外汇贵金属"].append("美元走强，非美货币和黄金承压")
        if "美元" in all_text and ("走弱" in all_text or "下跌" in all_text or "跌" in all_text):
            insights["外汇贵金属"].append("美元走弱，利好黄金和新兴市场货币")
        if "黄金" in all_text or "gold" in all_text:
            if "上涨" in all_text or "涨" in all_text:
                insights["外汇贵金属"].append("黄金上涨反映避险需求，关注地缘风险")
            elif "下跌" in all_text or "跌" in all_text:
                insights["外汇贵金属"].append("黄金回调，或受美元及实际利率影响")
            else:
                insights["外汇贵金属"].append("黄金市场波动，建议关注避险情绪变化")
        if "美债收益率" in all_text:
            insights["外汇贵金属"].append("美债收益率变动影响全球资产定价")
        if not insights["外汇贵金属"]:
            insights["外汇贵金属"].append("外汇市场波动正常，建议关注主要货币对走势")

    # 能源与大宗观点
    elif category == "能源与大宗":
        if "原油" in all_text or "油价" in all_text or "wti" in all_text or "布伦特" in all_text:
            if "上涨" in all_text or "涨" in all_text:
                insights["能源与大宗"].append("油价上涨，关注供应端变化及地缘因素")
            elif "下跌" in all_text or "跌" in all_text:
                insights["能源与大宗"].append("油价回落，或反映需求预期走弱")
            else:
                insights["能源与大宗"].append("原油市场关注度高，需跟踪OPEC+动向")
        if "opec" in all_text:
            insights["能源与大宗"].append("OPEC+政策是油价关键变量")
        if "天然气" in all_text:
            insights["能源与大宗"].append("天然气价格受季节性和供应因素影响")
        if "铜" in all_text or "铝" in all_text or "铁矿" in all_text:
            insights["能源与大宗"].append("工业金属价格反映制造业景气预期")
        if not insights["能源与大宗"]:
            insights["能源与大宗"].append("大宗商品整体平稳，关注供需基本面变化")

    # 股市公司观点
    elif category == "股市公司":
        if "财报" in all_text or "业绩" in all_text or "earnings" in all_text:
            insights["股市公司"].append("财报季期间，业绩表现将主导个股走势")
        if "回购" in all_text or "分红" in all_text:
            insights["股市公司"].append("公司回购/分红显示管理层对前景有信心")
        if "裁员" in all_text or "layoff" in all_text:
            insights["股市公司"].append("裁员消息或反映企业削减成本、调整预期")
        if "收购" in all_text or "并购" in all_text or "merger" in all_text:
            insights["股市公司"].append("并购活动活跃，关注行业整合机会")
        if "美股" in all_text or "纳指" in all_text or "标普" in all_text:
            insights["股市公司"].append("美股走势影响全球风险偏好")
        if "a股" in all_text or "港股" in all_text:
            insights["股市公司"].append("关注中国市场政策支持和资金流向")
        if not insights["股市公司"]:
            insights["股市公司"].append("股市波动正常，建议关注龙头企业动态")

    # 加密市场观点
    elif category == "加密市场":
        if "比特币" in all_text or "btc" in all_text:
            if "上涨" in all_text or "涨" in all_text or "新高" in all_text:
                insights["加密市场"].append("比特币走强，市场情绪偏乐观")
            elif "下跌" in all_text or "跌" in all_text:
                insights["加密市场"].append("比特币回调，注意市场波动风险")
            else:
                insights["加密市场"].append("比特币为加密市场风向标，需持续关注")
        if "以太坊" in all_text or "eth" in all_text:
            insights["加密市场"].append("以太坊生态发展值得关注")
        if "监管" in all_text or "regulation" in all_text:
            insights["加密市场"].append("监管政策是加密市场重要变量")
        if "交易所" in all_text:
            insights["加密市场"].append("交易所动态影响市场流动性和信心")
        if not insights["加密市场"]:
            insights["加密市场"].append("加密市场波动较大，建议谨慎参与")

    else:
        insights["其他"].append("其他资讯供参考，建议结合具体情况分析")

    # 返回最多3条观点
    result = insights.get(category, ["暂无特别观点"])
    return result[:3]


def format_daily_report_text(data: List[Dict[str, Any]], generated_at: str, title: str) -> str:
    grouped: Dict[str, List[Dict[str, Any]]] = {
        "宏观政策": [],
        "外汇贵金属": [],
        "能源与大宗": [],
        "股市公司": [],
        "加密市场": [],
        "其他": [],
    }

    sorted_data = sorted(data, key=lambda x: x.get("datetime", ""), reverse=True)
    for item in sorted_data:
        category = classify_news(item.get("content", ""))
        grouped[category].append(item)

    lines = []
    lines.append(f"生成时间: {generated_at}")
    lines.append(f"新闻总数: {len(data)}")

    if not data:
        lines.append("")
        lines.append("当日暂无符合条件的新闻。")
        lines.append("")
        lines.append("数据来源: Jin10")
        return "\n".join(lines)

    for category in ["宏观政策", "外汇贵金属", "能源与大宗", "股市公司", "加密市场", "其他"]:
        items = grouped[category]
        if not items:
            continue

        lines.append("")
        lines.append(f"【{category}】共 {len(items)} 条")

        # 核心要点
        lines.append("")
        lines.append("▎核心要点")
        key_points = extract_key_points(items, max_points=5)
        for idx, point in enumerate(key_points, 1):
            lines.append(f"  {idx}. {point}")

        # 观点分析
        lines.append("")
        lines.append("▎观点分析")
        insight_list = generate_category_insight(category, items)
        for insight in insight_list:
            lines.append(f"  • {insight}")

    lines.append("")
    lines.append("=" * 50)
    lines.append("免责声明: 以上内容仅供参考，不构成投资建议。")
    lines.append("数据来源: Jin10")
    return "\n".join(lines)


def _set_run_font(run, font_size=Pt(10.5), bold=False, color=None):
    """设置 run 字体：中文楷体，英文 Times New Roman"""
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color


def save_word_report(path: str, title: str, report_text: str) -> None:
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    _set_run_font(title_run, font_size=Pt(18), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    # Body
    for line in report_text.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)

        # 分类标题 【宏观政策】
        if line.startswith("【") and "】" in line:
            _set_run_font(run, font_size=Pt(14), bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        # 子标题 ▎核心要点 / ▎观点分析
        elif line.startswith("▎"):
            _set_run_font(run, font_size=Pt(11), bold=True, color=RGBColor(0x2E, 0x75, 0x2E))
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        # 编号列表   1. xxx
        elif re.match(r"^\s+\d+\.\s", line):
            _set_run_font(run, font_size=Pt(10.5))
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        # 观点列表（• 开头）
        elif "•" in line:
            _set_run_font(run, font_size=Pt(10.5), bold=True, color=RGBColor(0x8B, 0x45, 0x13))
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        # 元信息
        elif line.startswith("生成时间:") or line.startswith("新闻总数:") or line.startswith("数据来源:"):
            _set_run_font(run, font_size=Pt(10), color=RGBColor(0x66, 0x66, 0x66))
        # 免责声明
        elif "免责声明" in line:
            _set_run_font(run, font_size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))
            p.paragraph_format.space_before = Pt(8)
        # 分隔线
        elif line.startswith("="):
            _set_run_font(run, font_size=Pt(8), color=RGBColor(0xCC, 0xCC, 0xCC))
        else:
            _set_run_font(run, font_size=Pt(10.5))

        p.paragraph_format.line_spacing = 1.4

    doc.save(path)


def split_today_and_yesterday(data: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    today = datetime.now().date()
    yesterday = today - timedelta(days=1)

    today_items: List[Dict[str, Any]] = []
    yesterday_items: List[Dict[str, Any]] = []

    for item in data:
        dt_text = item.get("datetime", "")
        dt = None
        if dt_text:
            try:
                dt = datetime.strptime(dt_text, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                dt = None

        if not dt:
            raw_time = item.get("time", "")
            dt = parse_jin10_time(raw_time)

        if not dt:
            continue

        d = dt.date()
        if d == today:
            today_items.append(item)
        elif d == yesterday:
            yesterday_items.append(item)

    return {
        "today": sorted(today_items, key=lambda x: x.get("datetime", ""), reverse=True),
        "yesterday": sorted(yesterday_items, key=lambda x: x.get("datetime", ""), reverse=True),
    }


def save_two_word_reports(output_dir: str, data: List[Dict[str, Any]]) -> Dict[str, str]:
    """直接保存到指定目录，只生成两个 Word 文件"""
    os.makedirs(output_dir, exist_ok=True)

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    day_tag = datetime.now().strftime("%Y%m%d")
    split_result = split_today_and_yesterday(data)

    yesterday_title = "金十新闻摘要日报 - 昨日总结"
    today_title = "金十新闻摘要日报 - 今日最新"

    yesterday_items = split_result["yesterday"]
    if yesterday_items:
        yesterday_text = format_daily_report_text(yesterday_items, generated_at, yesterday_title)
    else:
        yesterday_text = "生成时间: {}\n新闻总数: 0\n\n当日暂无符合条件的新闻。".format(generated_at)

    today_text = format_daily_report_text(split_result["today"], generated_at, today_title)

    yesterday_docx = os.path.join(output_dir, f"昨日总结_{day_tag}.docx")
    today_docx = os.path.join(output_dir, f"今日最新_{day_tag}.docx")

    save_word_report(yesterday_docx, yesterday_title, yesterday_text)
    save_word_report(today_docx, today_title, today_text)

    return {
        "yesterday_docx": yesterday_docx,
        "today_docx": today_docx,
    }


def merge_news(*news_lists: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    merged: List[Dict[str, Any]] = []
    seen = set()

    for news_list in news_lists:
        for item in news_list:
            dedup_key = item.get("id") or f"{item.get('time', '')}|{item.get('content', '')[:60]}"
            if dedup_key in seen:
                continue
            seen.add(dedup_key)
            merged.append(item)

    return merged


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="抓取金十数据新闻/快讯")
    parser.add_argument("--pages", type=int, default=12, help="抓取页数，默认 12")
    parser.add_argument("--timeout", type=int, default=15, help="HTTP 超时秒数，默认 15")
    parser.add_argument("--delay", type=float, default=0.25, help="分页请求间隔秒数，默认 0.25")
    parser.add_argument("--since-hours", type=int, default=72, help="仅保留最近 N 小时，默认 72")
    parser.add_argument("--keyword", type=str, default="", help="关键词过滤，例如 美联储")
    parser.add_argument("--output-dir", type=str, default=".", help="输出目录，默认当前目录")
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)
    try:
        data = collect_news(
            pages=max(args.pages, 1),
            timeout=max(args.timeout, 5),
            delay=max(args.delay, 0.0),
            since_hours=args.since_hours if args.since_hours is None or args.since_hours >= 0 else None,
            keyword=args.keyword.strip() or None,
        )
    except requests.HTTPError as exc:
        print(f"[ERROR] HTTP 请求失败: {exc}")
        return 1
    except requests.RequestException as exc:
        print(f"[ERROR] 网络请求失败: {exc}")
        return 1
    except (ValueError, KeyError, json.JSONDecodeError) as exc:
        print(f"[ERROR] 响应解析失败: {exc}")
        return 1

    if not data:
        print("[WARN] 未抓取到满足条件的新闻。")
        return 0

    yesterday_date = datetime.now().date() - timedelta(days=1)
    try:
        yesterday_data = collect_news_for_date(
            target_date=yesterday_date,
            pages=max(args.pages, 8),
            timeout=max(args.timeout, 5),
            delay=max(args.delay, 0.0),
            keyword=args.keyword.strip() or None,
        )
        if yesterday_data:
            data = merge_news(data, yesterday_data)
            print(f"[INFO] 已补充昨日新闻 {len(yesterday_data)} 条")
        else:
            print("[WARN] 未补抓到昨日新闻，将使用归档兜底（若存在）")
    except Exception as exc:  # noqa: BLE001
        print(f"[WARN] 补抓昨日新闻失败，将使用归档兜底: {exc}")

    report_paths = save_two_word_reports(args.output_dir, data)

    print(f"[OK] 抓取完成，共 {len(data)} 条")
    print(f"[OK] 昨日总结(Word): {report_paths['yesterday_docx']}")
    print(f"[OK] 今日最新(Word): {report_paths['today_docx']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
